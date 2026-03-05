"""
Keiths Guitar Tab Converter
============================================
Upload a guitar tab PDF and receive a Keiths Guitar Tab .docx:
  • 14pt Times New Roman Bold throughout
  • 1.3 line spacing
  • Consecutive chord/lyric pairs merged side-by-side
  • 5–6 word lyric cues (extended, not full lyrics)
  • Artist name, Tuning, and Song Notes placed after Chord Diagrams
  • Chord diagrams for non-standard chords only

Requirements (install before running):
  pip install streamlit anthropic python-docx

Run:
  streamlit run app.py

You must set your Anthropic API key as an environment variable:
  export ANTHROPIC_API_KEY=your_key_here
"""

import io
import json
import os
import base64
import re
import streamlit as st
import anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Page config ────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Keith's Guitar Tab Converter",
    page_icon="🎸",
    layout="centered",
)

# ─── Styles ─────────────────────────────────────────────────────────────────

st.markdown("""
<style>
    .main-title { font-size: 2rem; font-weight: 700; color: #1a1a2e; margin-bottom: 0; }
    .subtitle   { font-size: 1rem; color: #555; margin-top: 0.2rem; margin-bottom: 1.5rem; }
    .rule-box   { background: #f8f9fa; border-left: 4px solid #4a90d9;
                  padding: 0.8rem 1rem; border-radius: 4px; font-size: 0.9rem; }
    .success-box { background: #e8f5e9; border-left: 4px solid #43a047;
                   padding: 0.8rem 1rem; border-radius: 4px; }
</style>
""", unsafe_allow_html=True)


# ─── Header ─────────────────────────────────────────────────────────────────

st.markdown('<p class="main-title">🎸 Keiths Tab Maker</p>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Guitar Tab Converter — Make Charts Like Marcs</p>', unsafe_allow_html=True)

with st.expander("ℹ️ What does this converter do?"):
    st.markdown("""
**Keiths Tab Maker** converts guitar tab PDFs into compact `.docx` chord charts optimised for live performance:

| Feature | Value |
|---|---|
| Font | Times New Roman Bold, 14pt |
| Line spacing | 1.3 |
| Lyric mode | Extended cues (5–6 words) **or** your own full lyrics |
| Line merging | Consecutive chord/lyric pairs merged side-by-side |
| Chord notation | `ChordName-Beats` (e.g. `Am-4`, `G7sus4-2`) |
| Non-standard chord diagrams | Included at bottom |
| Artist / Tuning / Notes | After chord diagrams |

The side-by-side merging reduces vertical space by ~50% vs. a standard chart, while keeping chords and cues clearly readable.

**Lyric modes:**
- **Extended cues** — Claude automatically generates the first 5–6 words of each line as a memory prompt
- **Full lyrics** — paste the complete lyrics from your own licensed source and they will be placed exactly as written
""")


# ─── API key ────────────────────────────────────────────────────────────────

api_key = os.environ.get("ANTHROPIC_API_KEY", "")
if not api_key:
    api_key = st.text_input(
        "Anthropic API Key",
        type="password",
        placeholder="sk-ant-...",
        help="Your key is used only for this request and never stored.",
    )

if not api_key:
    st.info("Enter your Anthropic API key above to get started.")
    st.stop()


# ─── File upload ────────────────────────────────────────────────────────────

uploaded = st.file_uploader(
    "Upload a guitar tab PDF",
    type=["pdf"],
    help="Upload a chord chart or guitar tab PDF from any source.",
)

if not uploaded:
    st.stop()

st.success(f"✅ Uploaded: **{uploaded.name}**")

# ─── Lyric mode selector ─────────────────────────────────────────────────────

st.markdown("#### 🎤 Lyric Mode")
lyric_mode = st.radio(
    "How should lyrics appear in your chart?",
    options=["Extended cues (5–6 words, auto-generated)", "Full lyrics (paste your own)"],
    index=0,
    help="Extended cues are memory prompts Claude generates automatically. Full lyrics requires you to paste the complete text."
)

user_lyrics = ""
if lyric_mode == "Full lyrics (paste your own)":
    st.info("Paste the complete lyrics below, section by section, in the same order they appear in the song. Use blank lines to separate sections.")
    user_lyrics = st.text_area(
        "Paste full lyrics here",
        height=300,
        placeholder="[Verse 1]\nWell, sometimes I go out by myself\nAnd I look across the water\n...\n\n[Chorus]\nValerie\nValerie\n...",
        help="Paste lyrics in order, with section labels matching the tab (Verse 1, Chorus, Bridge, etc.)"
    )
    if not user_lyrics.strip():
        st.warning("Please paste your lyrics above before converting.")



# ─── DOCX builder helpers ────────────────────────────────────────────────────

FONT_NAME  = "Times New Roman"
FONT_SIZE  = Pt(14)
LINE_RULE  = 374        # 1.3 × 288 twips ≈ 374 (WD_LINE_SPACING.MULTIPLE not needed; direct twip value)
TAB_STOP   = Inches(3.4)


def _apply_run(run):
    run.font.name    = FONT_NAME
    run.font.size    = FONT_SIZE
    run.font.bold    = True
    # Ensure East-Asian and complex-script fonts match
    rpr = run._r.get_or_add_rPr()
    for tag in ("w:rFonts",):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.insert(0, el)
        el.set(qn("w:ascii"),    FONT_NAME)
        el.set(qn("w:hAnsi"),    FONT_NAME)
        el.set(qn("w:cs"),       FONT_NAME)


def _set_spacing(para, before=0, after=80, line=LINE_RULE):
    pPr  = para._p.get_or_add_pPr()
    spc  = pPr.find(qn("w:spacing"))
    if spc is None:
        spc = OxmlElement("w:spacing")
        pPr.append(spc)
    spc.set(qn("w:before"),   str(before))
    spc.set(qn("w:after"),    str(after))
    spc.set(qn("w:line"),     str(line))
    spc.set(qn("w:lineRule"), "auto")


def _add_separator(doc):
    para = doc.add_paragraph()
    _set_spacing(para, before=60, after=60)
    run = para.add_run("")
    _apply_run(run)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_label(doc, text):
    para = doc.add_paragraph()
    _set_spacing(para, before=140, after=0)
    run = para.add_run(text)
    _apply_run(run)


def _add_line(doc, text, before=100, after=80):
    para = doc.add_paragraph()
    _set_spacing(para, before=before, after=after)
    run = para.add_run(text)
    _apply_run(run)


def _add_tab_stop(para, position):
    pPr  = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(position.twips)))
    tabs.append(tab)


def _add_merged_chord(doc, left, right):
    para = doc.add_paragraph()
    _set_spacing(para, before=100, after=0)
    _add_tab_stop(para, TAB_STOP)
    r1 = para.add_run(left + "  |  ")
    _apply_run(r1)
    r_tab = para.add_run("\t")
    _apply_run(r_tab)
    r2 = para.add_run(right)
    _apply_run(r2)


def _add_merged_lyric(doc, left, right):
    para = doc.add_paragraph()
    _set_spacing(para, before=0, after=80)
    _add_tab_stop(para, TAB_STOP)
    r1 = para.add_run(left)
    _apply_run(r1)
    r_tab = para.add_run("\t")
    _apply_run(r_tab)
    r2 = para.add_run(right)
    _apply_run(r2)


def _add_blank(doc):
    para = doc.add_paragraph()
    _set_spacing(para, before=0, after=0)
    run = para.add_run("")
    _apply_run(run)


def build_docx(song: dict) -> bytes:
    """
    Build the Keiths Tab .docx from a structured song dict.

    Expected song dict shape:
    {
      "title": str,
      "sections": [
        {
          "label": str,            # e.g. "[Verse 1]", "CHORUS", "PRE-CHORUS"
          "type": str,             # "lyrics" | "instrumental"
          "pairs": [               # for lyrics sections
            {"chords": str, "lyric": str}
          ],
          "bars": str              # for instrumental sections (pipe notation)
        }
      ],
      "chord_diagrams": [
        {"name": str, "voicing": str, "notes": str}
      ],
      "artist": str,
      "tuning": str,
      "song_notes": [str]
    }
    """
    doc = Document()

    # Page size & margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin  = section.bottom_margin = Inches(1)

    # Title (centred)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(title_para, before=0, after=200)
    t_run = title_para.add_run(song.get("title", "UNTITLED").upper())
    _apply_run(t_run)

    # Sections
    for sec in song.get("sections", []):
        _add_separator(doc)
        _add_label(doc, sec["label"])

        if sec.get("type") == "instrumental":
            _add_line(doc, sec.get("bars", ""), before=100, after=80)
            _add_blank(doc)
            continue

        pairs = sec.get("pairs", [])
        i = 0
        while i < len(pairs):
            if i + 1 < len(pairs):
                # Merge pair i and pair i+1 side-by-side
                p1 = pairs[i]
                p2 = pairs[i + 1]
                _add_merged_chord(doc, p1["chords"], p2["chords"])
                _add_merged_lyric(doc, p1["lyric"],  p2["lyric"])
                i += 2
            else:
                # Odd one out — stays on its own line
                p = pairs[i]
                _add_line(doc, p["chords"], before=100, after=0)
                _add_line(doc, p["lyric"],  before=0,   after=80)
                i += 1

        _add_blank(doc)

    # Chord diagrams
    diagrams = song.get("chord_diagrams", [])
    if diagrams:
        _add_separator(doc)
        _add_label(doc, "CHORD DIAGRAMS")
        _add_blank(doc)
        for d in diagrams:
            _add_line(doc, f"{d['name']}  —  {d.get('notes', '')}", before=100, after=0)
            _add_line(doc, d.get("voicing", ""), before=0, after=100)

    # Song info (artist / tuning / notes)
    _add_blank(doc)
    _add_separator(doc)

    if song.get("artist"):
        _add_label(doc, "ARTIST")
        _add_line(doc, song["artist"], before=60, after=60)
        _add_blank(doc)

    if song.get("tuning"):
        _add_label(doc, "TUNING")
        _add_line(doc, song["tuning"], before=60, after=60)
        _add_blank(doc)

    notes = song.get("song_notes", [])
    if notes:
        _add_label(doc, "SONG NOTES")
        for note in notes:
            _add_line(doc, note, before=60, after=60)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Claude prompts ──────────────────────────────────────────────────────────

SYSTEM_PROMPT_CUES = """You are an expert guitar tab formatter. Your job is to parse a guitar tab PDF and return a structured JSON object that will be used to build a Keiths Tab Extended chord chart.

RULES FOR PARSING:
1. Extract the song title, artist, tuning/key, and any performance notes.
2. Identify every section: Intro, Verse 1/2/3, Pre-Chorus, Chorus, Bridge, Solo, Outro, Instrumental, Coda, etc.
3. For each LYRIC section, produce an array of "pairs" — each pair has:
   - "chords": the chord line with ChordName-Beats notation (e.g. "Am-4", "G7sus4-2")
     Position chords using spaces so they align above the correct lyric word.
   - "lyric": the first 5-6 words of that lyric line followed by "..."
     Example: "Well, sometimes, I go out..." (NOT the full lyric)
4. For INSTRUMENTAL sections (Intro, Solo, Instrumental breaks), use pipe bar notation:
   "|A-4|D-2|E-2|" — set type to "instrumental" and put this in "bars".
5. Solo sections must include placement context in the label:
   "SOLO — appears after the second chorus"
6. Chord diagrams: include ONLY non-standard chords (barre chords, maj7, sus, dim, aug, 7th chords with unusual voicings). DO NOT include diagrams for: C, D, E, G, A, Am, Em, Dm, F.
7. Beat counts: if not shown in the source, infer from the time signature and lyric density. A full bar in 4/4 = 4 beats. Common splits: one chord per bar = -4, two chords = -2 each.
8. N.C. = No Chord, keep as "N.C.-beats".

OUTPUT FORMAT — return ONLY valid JSON, no markdown fences, no explanation:
{
  "title": "SONG TITLE",
  "artist": "Artist Name",
  "tuning": "E A D G B E (standard) | Key: X | Tempo: N bpm",
  "sections": [
    {
      "label": "[Verse 1]",
      "type": "lyrics",
      "pairs": [
        {"chords": "  Am-4          G-4", "lyric": "First five or six words..."},
        {"chords": "  F-2    C-2", "lyric": "Next five or six words..."}
      ]
    },
    {
      "label": "INTRO",
      "type": "instrumental",
      "bars": "|A-4|A-4|D-2|E-2|  2X"
    }
  ],
  "chord_diagrams": [
    {
      "name": "Bsus4",
      "voicing": "x  2  4  4  4  2    Strings: E A D G B e  (barre 2fr)    fingers: 1-3-4",
      "notes": "2fr barre"
    }
  ],
  "song_notes": [
    "Any capo instructions, FILL descriptions, strumming notes, etc."
  ]
}"""


SYSTEM_PROMPT_FULL_LYRICS = """You are an expert guitar tab formatter. Your job is to parse a guitar tab PDF and a set of user-provided lyrics, then return a structured JSON object for a Keiths Tab Extended chord chart.

RULES FOR PARSING:
1. Extract the song title, artist, tuning/key, and any performance notes from the PDF.
2. Identify every section: Intro, Verse 1/2/3, Pre-Chorus, Chorus, Bridge, Solo, Outro, Instrumental, Coda, etc.
3. For each LYRIC section, produce an array of "pairs". Match each chord line from the PDF to its corresponding lyric line from the user-supplied lyrics. Each pair has:
   - "chords": the chord line with ChordName-Beats notation (e.g. "Am-4", "G7sus4-2")
     Position chords using spaces so they align above the correct word in the lyric line.
   - "lyric": the COMPLETE lyric line exactly as the user provided it — do not truncate, do not add "..."
4. For INSTRUMENTAL sections (Intro, Solo, Instrumental breaks), use pipe bar notation:
   "|A-4|D-2|E-2|" — set type to "instrumental" and put this in "bars".
5. Solo sections must include placement context in the label:
   "SOLO — appears after the second chorus"
6. Chord diagrams: include ONLY non-standard chords (barre chords, maj7, sus, dim, aug, 7th chords with unusual voicings). DO NOT include diagrams for: C, D, E, G, A, Am, Em, Dm, F.
7. Beat counts: if not shown in the source, infer from the time signature and lyric density.
8. N.C. = No Chord, keep as "N.C.-beats".
9. If the user has provided lyrics for a section that repeats (e.g. the chorus appears 3 times), use the provided lyrics for each occurrence — do not skip repeated sections.

OUTPUT FORMAT — return ONLY valid JSON, no markdown fences, no explanation:
{
  "title": "SONG TITLE",
  "artist": "Artist Name",
  "tuning": "E A D G B E (standard) | Key: X | Tempo: N bpm",
  "sections": [
    {
      "label": "[Verse 1]",
      "type": "lyrics",
      "pairs": [
        {"chords": "  Am-4          G-4", "lyric": "Complete lyric line exactly as provided"},
        {"chords": "  F-2    C-2", "lyric": "Next complete lyric line exactly as provided"}
      ]
    },
    {
      "label": "INTRO",
      "type": "instrumental",
      "bars": "|A-4|A-4|D-2|E-2|  2X"
    }
  ],
  "chord_diagrams": [
    {
      "name": "Bsus4",
      "voicing": "x  2  4  4  4  2    Strings: E A D G B e  (barre 2fr)    fingers: 1-3-4",
      "notes": "2fr barre"
    }
  ],
  "song_notes": [
    "Any capo instructions, FILL descriptions, strumming notes, etc."
  ]
}"""


USER_PROMPT_CUES = """Please parse this guitar tab PDF and return the structured JSON as described.
Remember:
- Lyric cues = first 5-6 words only, followed by "..."
- Chord notation = ChordName-Beats (e.g. Ebmaj7-4, Fm-2)
- Solo labels must include placement context
- Only non-standard chord diagrams
- Return ONLY the JSON object, nothing else"""


def build_user_prompt_full_lyrics(lyrics: str) -> str:
    return f"""Please parse this guitar tab PDF together with the user-supplied lyrics below, and return the structured JSON as described.

USER-SUPPLIED LYRICS:
{lyrics}

Remember:
- Use the COMPLETE lyric lines exactly as provided above — do not truncate them
- Match each lyric line to its corresponding chord line from the PDF
- Chord notation = ChordName-Beats (e.g. Ebmaj7-4, Fm-2)
- Solo labels must include placement context
- Only non-standard chord diagrams
- Return ONLY the JSON object, nothing else"""


# ─── Conversion logic ────────────────────────────────────────────────────────

def parse_tab_with_claude(pdf_bytes: bytes, api_key: str, system_prompt: str, user_prompt: str) -> dict:
    client   = anthropic.Anthropic(api_key=api_key)
    b64_pdf  = base64.standard_b64encode(pdf_bytes).decode("utf-8")

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        system=system_prompt,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": b64_pdf,
                        },
                    },
                    {"type": "text", "text": user_prompt},
                ],
            }
        ],
    )

    raw = message.content[0].text.strip()
    # Strip any accidental markdown fences
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"```\s*$", "",  raw)
    return json.loads(raw)


# ─── Main convert button ─────────────────────────────────────────────────────

# Block conversion if full lyrics mode selected but no lyrics pasted
full_lyrics_mode = lyric_mode == "Full lyrics (paste your own)"
ready_to_convert = not full_lyrics_mode or user_lyrics.strip()

col1, col2 = st.columns([2, 1])
with col1:
    convert_btn = st.button(
        "🎵 Convert to Keiths Tab Extended",
        type="primary",
        use_container_width=True,
        disabled=not ready_to_convert
    )
with col2:
    st.markdown("")   # spacer

if convert_btn:
    pdf_bytes = uploaded.read()

    # Select prompts based on lyric mode
    if full_lyrics_mode:
        system_prompt = SYSTEM_PROMPT_FULL_LYRICS
        user_prompt   = build_user_prompt_full_lyrics(user_lyrics)
        mode_label    = "Full lyrics"
    else:
        system_prompt = SYSTEM_PROMPT_CUES
        user_prompt   = USER_PROMPT_CUES
        mode_label    = "Extended cues"

    with st.spinner(f"Parsing tab with Claude ({mode_label} mode)..."):
        try:
            song_data = parse_tab_with_claude(pdf_bytes, api_key, system_prompt, user_prompt)
        except json.JSONDecodeError as e:
            st.error(f"Could not parse Claude's response as JSON: {e}")
            st.stop()
        except anthropic.APIError as e:
            st.error(f"Anthropic API error: {e}")
            st.stop()

    with st.spinner("Building .docx..."):
        docx_bytes = build_docx(song_data)

    title_slug = re.sub(r"[^a-zA-Z0-9]+", "_", song_data.get("title", "song")).lower()
    filename   = f"{title_slug}_space_saver_extended.docx"

    st.markdown('<div class="success-box">✅ Conversion complete! Download your chart below.</div>',
                unsafe_allow_html=True)
    st.markdown("")

    st.download_button(
        label     = f"⬇️ Download {filename}",
        data      = docx_bytes,
        file_name = filename,
        mime      = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    # Preview the parsed structure
    with st.expander("🔍 Preview parsed song structure"):
        st.write(f"**Title:** {song_data.get('title')}")
        st.write(f"**Artist:** {song_data.get('artist')}")
        st.write(f"**Tuning:** {song_data.get('tuning')}")
        st.write(f"**Lyric mode used:** {mode_label}")
        st.write(f"**Sections:** {len(song_data.get('sections', []))}")
        for sec in song_data.get("sections", []):
            pairs = sec.get("pairs", [])
            bars  = sec.get("bars", "")
            if pairs:
                st.write(f"  • **{sec['label']}** — {len(pairs)} chord/lyric pairs → {len(pairs)//2 + len(pairs)%2} merged rows")
            else:
                st.write(f"  • **{sec['label']}** (instrumental) — `{bars}`")
        diagrams = song_data.get("chord_diagrams", [])
        if diagrams:
            st.write(f"**Chord diagrams:** {', '.join(d['name'] for d in diagrams)}")


# ─── Footer ─────────────────────────────────────────────────────────────────

st.markdown("---")
st.markdown(
    "<small>Keiths Tab Extended • 14pt Times New Roman Bold • 1.3 spacing • "
    "Side-by-side line merging • Extended cues or full lyrics • Song info after diagrams</small>",
    unsafe_allow_html=True,
)